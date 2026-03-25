#!/usr/bin/env python3
"""
MD → DOCX 변환 스크립트 — 서식2 (도전적 문제 정의서)
원본 HWPX 양식의 표 레이아웃을 Word로 재현:
  - 도전적 문제 표 (2행×3열, 라벨 rowSpan=2)
  - 미션/제안자 표 (3행×6열, 복합 병합)
  - 본문 대형 표 (10행×2열, 라벨+내용 쌍)
  - 참고자료 행 (1행)

색상: 라벨 #DFEAF5 (연파랑), 헤더 #2B5686 (진파랑), 흰글씨
폰트: 맑은 고딕 10pt, 줄간격 130%, Justify
여백: 상하 15mm, 좌우 20mm
"""
import re
import sys
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ── 색상 상수 (서식2 원본 기준) ──
COLOR_HEADER_BG = "2B5686"    # 진파랑 배경 (헤더 셀)
COLOR_LABEL_BG = "DFEAF5"     # 연파랑 배경 (라벨 셀)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)
COLOR_BLUE_LINK = RGBColor(0x04, 0x32, 0xFF)

# 표 너비 (dxa = 1/20 pt). 서식2 원본 본문폭 ≈ 170mm
TABLE_WIDTH_DXA = 9639  # ≈ 170mm


def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_vertical_align(cell, val="center"):
    tc_pr = cell._tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{val}"/>')
    tc_pr.append(vAlign)


def set_cell_width(cell, width_dxa):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width_dxa}" w:type="dxa"/>')
    tcPr.append(tcW)


def set_cell_margins(cell, top=57, bottom=57, left=57, right=57):
    """셀 여백 설정 (dxa 단위)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(margins)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """셀 테두리 개별 설정. 값은 (sz, color) 튜플 또는 None"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    parts = []
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            sz, color = val
            parts.append(f'<w:{side} w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>')
        else:
            parts.append(f'<w:{side} w:val="single" w:sz="2" w:space="0" w:color="999999"/>')
    borders_xml = f'<w:tcBorders {nsdecls("w")}>{"".join(parts)}</w:tcBorders>'
    tcPr.append(parse_xml(borders_xml))


def set_paragraph_spacing(para, before=0, after=0, line=260, line_rule='auto'):
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = parse_xml(f'<w:spacing {nsdecls("w")}/>')
        pPr.append(spacing)
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    spacing.set(qn('w:line'), str(line))
    spacing.set(qn('w:lineRule'), line_rule)


def add_run(para, text, bold=False, italic=False, size=None, color=None, font_name='맑은 고딕', spacing=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    if spacing is not None:
        sp = parse_xml(f'<w:spacing {nsdecls("w")} w:val="{spacing}"/>')
        rPr.append(sp)
    return run


def add_formatted_text(para, text, base_bold=False, base_size=10, base_color=None, font_name='맑은 고딕', spacing=None):
    """Markdown 인라인 서식 파싱"""
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))'
    for match in re.finditer(pattern, text):
        if match.group(2):
            add_run(para, match.group(2), bold=True, italic=True, size=base_size, color=base_color, font_name=font_name, spacing=spacing)
        elif match.group(3):
            add_run(para, match.group(3), bold=True, size=base_size, color=base_color, font_name=font_name, spacing=spacing)
        elif match.group(4):
            add_run(para, match.group(4), bold=base_bold, italic=True, size=base_size, color=base_color, font_name=font_name, spacing=spacing)
        elif match.group(5):
            add_run(para, match.group(5), bold=base_bold, size=base_size, color=base_color, font_name=font_name, spacing=spacing)


def set_table_borders(table, outer_sz=8, outer_color="auto", inner_sz=2, inner_color="999999"):
    """표 전체 테두리"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{outer_sz}" w:space="0" w:color="{outer_color}"/>'
        f'  <w:left w:val="single" w:sz="{outer_sz}" w:space="0" w:color="{outer_color}"/>'
        f'  <w:bottom w:val="single" w:sz="{outer_sz}" w:space="0" w:color="{outer_color}"/>'
        f'  <w:right w:val="single" w:sz="{outer_sz}" w:space="0" w:color="{outer_color}"/>'
        f'  <w:insideH w:val="single" w:sz="{inner_sz}" w:space="0" w:color="{inner_color}"/>'
        f'  <w:insideV w:val="single" w:sz="{inner_sz}" w:space="0" w:color="{inner_color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def set_table_width(table, width_dxa):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{width_dxa}" w:type="dxa"/>')
    tblPr.append(tblW)


# ────────────────────────────────────────────
# 표1: 도전적 문제 (2행×3열)
# ────────────────────────────────────────────
def create_problem_table(doc, title_kr, title_en):
    """도전적 문제 표: 좌측 라벨(2행 병합) + 국문/영문"""
    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, TABLE_WIDTH_DXA)
    set_table_borders(table, outer_sz=8, outer_color="auto", inner_sz=2, inner_color="999999")

    col_widths = [748, 1214, 7677]  # 비율 유지

    # 1열: "도전적\n문제" 2행 병합
    cell_a1 = table.cell(0, 0)
    cell_a2 = table.cell(1, 0)
    merged = cell_a1.merge(cell_a2)
    set_cell_shading(merged, COLOR_HEADER_BG)
    set_cell_vertical_align(merged, "center")
    set_cell_width(merged, col_widths[0])
    set_cell_margins(merged, top=57, bottom=57, left=57, right=57)

    p = merged.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line=260)
    add_run(p, "도전적", bold=True, size=10, color=COLOR_WHITE, spacing=-5)
    p2 = merged.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p2, line=260)
    add_run(p2, "문제", bold=True, size=10, color=COLOR_WHITE, spacing=-5)

    # 국문 행
    for row_idx, (label, content) in enumerate([("[국문]", title_kr), ("[영문]", title_en)]):
        label_cell = table.cell(row_idx, 1)
        set_cell_shading(label_cell, COLOR_LABEL_BG if row_idx == 0 else COLOR_LABEL_BG)
        set_cell_vertical_align(label_cell, "center")
        set_cell_width(label_cell, col_widths[1])
        lp = label_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(lp, line=260)
        add_run(lp, label, bold=True, size=10)

        content_cell = table.cell(row_idx, 2)
        set_cell_width(content_cell, col_widths[2])
        set_cell_vertical_align(content_cell, "center")
        cp = content_cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_spacing(cp, line=260)
        add_formatted_text(cp, content, base_bold=True, base_size=10)

    return table


# ────────────────────────────────────────────
# 표2: 미션 + 제안자 인적사항 (3행×6열)
# ────────────────────────────────────────────
def create_info_table(doc, mission_text, affiliation, name, contact, privacy_consent):
    """미션/제안자 인적사항 표"""
    table = doc.add_table(rows=3, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, TABLE_WIDTH_DXA)
    set_table_borders(table, outer_sz=8, outer_color="auto", inner_sz=2, inner_color="999999")

    # 열 너비: 라벨(748) + 5열
    col_w = [748, 1210, 3413, 962, 1480, 1826]

    # Row 0: 미션
    mission_label = table.cell(0, 0)
    set_cell_shading(mission_label, COLOR_HEADER_BG)
    set_cell_vertical_align(mission_label, "center")
    set_cell_width(mission_label, col_w[0])
    p = mission_label.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line=260)
    add_run(p, "미션", bold=True, size=10, color=COLOR_WHITE, spacing=-5)

    # 미션 내용: 5열 병합
    mission_content = table.cell(0, 1)
    for ci in range(2, 6):
        mission_content = mission_content.merge(table.cell(0, ci))
    p = mission_content.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, line=260)
    add_formatted_text(p, mission_text, base_size=10, spacing=-8)

    # Row 1-2: 제안자 인적사항 (라벨 2행 병합)
    proposer_label = table.cell(1, 0)
    proposer_label2 = table.cell(2, 0)
    merged_proposer = proposer_label.merge(proposer_label2)
    set_cell_shading(merged_proposer, COLOR_HEADER_BG)
    set_cell_vertical_align(merged_proposer, "center")
    set_cell_width(merged_proposer, col_w[0])
    set_cell_margins(merged_proposer, top=0, bottom=0, left=0, right=0)

    p = merged_proposer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line=220)
    add_run(p, "제안자", bold=True, size=10, color=COLOR_WHITE)
    p2 = merged_proposer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p2, line=220)
    add_run(p2, "인적", bold=True, size=10, color=COLOR_WHITE)
    p3 = merged_proposer.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p3, line=220)
    add_run(p3, "사항", bold=True, size=10, color=COLOR_WHITE)

    # Row 1: 소속 + 성함
    for ci, width in enumerate(col_w[1:], 1):
        set_cell_width(table.cell(1, ci), width)
        set_cell_width(table.cell(2, ci), width)

    # 소속 라벨
    aff_label = table.cell(1, 1)
    p = aff_label.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line=260)
    add_run(p, "소 속", bold=True, size=10, spacing=-7)
    set_cell_vertical_align(aff_label, "center")

    # 소속 내용
    aff_content = table.cell(1, 2)
    p = aff_content.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, line=260)
    add_run(p, affiliation, size=10, spacing=-7)
    set_cell_vertical_align(aff_content, "center")

    # 성함 라벨
    name_label = table.cell(1, 3)
    p = name_label.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line=260)
    add_run(p, "성 함", bold=True, size=10, spacing=-7)
    set_cell_vertical_align(name_label, "center")

    # 성함 내용 (2열 병합)
    name_content = table.cell(1, 4).merge(table.cell(1, 5))
    p = name_content.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line=260)
    add_run(p, name, size=10, spacing=-7)
    set_cell_vertical_align(name_content, "center")

    # Row 2: 연락처 + 개인정보 수집 동의
    contact_label = table.cell(2, 1)
    p = contact_label.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line=260)
    add_run(p, "연락처", bold=True, size=10, spacing=-7)
    set_cell_vertical_align(contact_label, "center")

    contact_content = table.cell(2, 2)
    p = contact_content.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, line=260)
    add_run(p, contact, size=10, spacing=-7)
    set_cell_vertical_align(contact_content, "center")

    # 개인정보 수집 동의 (2열 병합)
    privacy_label = table.cell(2, 3).merge(table.cell(2, 4))
    p = privacy_label.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line=260)
    add_run(p, "개인정보 수집 동의", bold=True, size=10, spacing=-7)
    set_cell_vertical_align(privacy_label, "center")

    privacy_content = table.cell(2, 5)
    p = privacy_content.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line=260)
    consent_text = "O / X" if privacy_consent == "O" else "O / X"
    # O에 파란색, X에 검정색 (또는 반대)
    if privacy_consent == "O":
        add_run(p, " O ", size=10, color=COLOR_BLUE_LINK, spacing=-7)
        add_run(p, "/ X ", size=10, spacing=-7)
    else:
        add_run(p, " O / ", size=10, spacing=-7)
        add_run(p, "X ", size=10, color=COLOR_BLUE_LINK, spacing=-7)

    return table


# ────────────────────────────────────────────
# 표3: 본문 대형 표 (섹션별 라벨+내용)
# ────────────────────────────────────────────
SECTION_LABELS = {
    "1": ("정의 및 중요성", "무엇이 문제인가?"),
    "2": ("2. 배경 및 목적", "왜 이 문제에 대한 해결을 시도하려고 하는가?"),
    "3": ("3. 목표 및 방법", "이 문제를 해결하기 위한 기존의 방법과 한계는 무엇이며, 시도하고자 하는 방법은?"),
    "4": ("4. 예상 결과", "문제를 해결할 수 있을 것이라는 신호는 무엇인가?"),
    "5": ("5. 파급효과", "만약 성공한다면, 그 영향은 무엇인가?"),
    "6": ("6. 참고자료", "각 질문에 대해 입증할 수 있는 학술 자료, 통계, 보고서, 실험/임상 데이터 등의 참고자료"),
}


def add_content_paragraphs(cell, content_lines, is_reference=False):
    """셀에 내용 문단들을 추가"""
    first = True
    for line_data in content_lines:
        line_type = line_data['type']
        text = line_data['text']

        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()

        if line_type == 'heading3':
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(p, before=80, after=40, line=286)
            add_formatted_text(p, text, base_bold=True, base_size=10, spacing=-10)
        elif line_type == 'heading4':
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(p, before=60, after=30, line=286)
            add_formatted_text(p, text, base_bold=True, base_size=10, spacing=-10)
        elif line_type == 'bullet':
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(p, before=20, after=20, line=286)
            pPr = p._p.get_or_add_pPr()
            ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="284" w:hanging="142"/>')
            pPr.append(ind)
            add_run(p, "- ", size=10, spacing=-10)
            add_formatted_text(p, text, base_size=10, spacing=-10)
        elif line_type == 'numbered':
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(p, before=20, after=20, line=286)
            pPr = p._p.get_or_add_pPr()
            ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="284" w:hanging="284"/>')
            pPr.append(ind)
            add_formatted_text(p, text, base_size=10, spacing=-10)
        elif line_type == 'table':
            # 표 데이터는 별도 처리
            pass
        elif line_type == 'reference':
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(p, before=20, after=20, line=260)
            add_formatted_text(p, text, base_size=10, spacing=-10)
        else:
            # 일반 본문
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(p, before=30, after=30, line=286)
            add_formatted_text(p, text, base_size=10, spacing=-10)

    if first:
        # 빈 셀
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, line=260)


def create_nested_table_in_cell(cell, rows_data):
    """셀 내부에 실제 Word 중첩 표를 생성"""
    if not rows_data or len(rows_data) < 2:
        return

    n_cols = len(rows_data[0])
    n_rows = len(rows_data)

    # 셀의 tc 요소에 직접 tbl을 삽입
    tc = cell._tc

    # 중첩 표 XML 생성
    tbl = parse_xml(f'<w:tbl {nsdecls("w")}/>')
    tblPr = parse_xml(
        f'<w:tblPr {nsdecls("w")}>'
        f'  <w:tblW w:w="0" w:type="auto"/>'
        f'  <w:tblBorders>'
        f'    <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        f'    <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        f'    <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        f'    <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        f'    <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        f'    <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        f'  </w:tblBorders>'
        f'  <w:tblLayout w:type="autofit"/>'
        f'  <w:tblCellMar>'
        f'    <w:top w:w="28" w:type="dxa"/>'
        f'    <w:left w:w="57" w:type="dxa"/>'
        f'    <w:bottom w:w="28" w:type="dxa"/>'
        f'    <w:right w:w="57" w:type="dxa"/>'
        f'  </w:tblCellMar>'
        f'</w:tblPr>'
    )
    tbl.append(tblPr)

    # tblGrid
    tblGrid = parse_xml(f'<w:tblGrid {nsdecls("w")}/>')
    for _ in range(n_cols):
        gridCol = parse_xml(f'<w:gridCol {nsdecls("w")} w:w="0"/>')
        tblGrid.append(gridCol)
    tbl.append(tblGrid)

    for row_idx, row_data in enumerate(rows_data):
        tr = parse_xml(f'<w:tr {nsdecls("w")}/>')
        is_header = (row_idx == 0)

        for col_idx in range(n_cols):
            tc_inner = parse_xml(f'<w:tc {nsdecls("w")}/>')
            tcPr_inner = parse_xml(f'<w:tcPr {nsdecls("w")}/>')

            # 헤더 행: 진파랑 배경 + 흰 글씨
            if is_header:
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{COLOR_HEADER_BG}" w:val="clear"/>')
                tcPr_inner.append(shd)

            tc_inner.append(tcPr_inner)

            # 셀 내용
            cell_text = row_data[col_idx].strip() if col_idx < len(row_data) else ""
            p_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
            pPr_xml = parse_xml(
                f'<w:pPr {nsdecls("w")}>'
                f'  <w:spacing w:before="0" w:after="0" w:line="220" w:lineRule="auto"/>'
                f'  <w:jc w:val="center"/>'
                f'</w:pPr>'
            )
            p_xml.append(pPr_xml)

            # 텍스트 run
            r_xml = parse_xml(f'<w:r {nsdecls("w")}/>')
            rPr_xml = parse_xml(f'<w:rPr {nsdecls("w")}/>')
            sz = parse_xml(f'<w:sz {nsdecls("w")} w:val="16"/>')  # 8pt
            szCs = parse_xml(f'<w:szCs {nsdecls("w")} w:val="16"/>')
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="맑은 고딕" w:eastAsia="맑은 고딕" w:hAnsi="맑은 고딕"/>')
            sp = parse_xml(f'<w:spacing {nsdecls("w")} w:val="-10"/>')
            rPr_xml.append(rFonts)
            rPr_xml.append(sz)
            rPr_xml.append(szCs)
            rPr_xml.append(sp)

            if is_header:
                bold_xml = parse_xml(f'<w:b {nsdecls("w")}/>')
                color_xml = parse_xml(f'<w:color {nsdecls("w")} w:val="FFFFFF"/>')
                rPr_xml.append(bold_xml)
                rPr_xml.append(color_xml)

            r_xml.append(rPr_xml)
            t_xml = parse_xml(f'<w:t {nsdecls("w")} xml:space="preserve"/>')
            t_xml.text = cell_text
            r_xml.append(t_xml)
            p_xml.append(r_xml)

            tc_inner.append(p_xml)
            tr.append(tc_inner)

        tbl.append(tr)

    # 표 뒤에 필수 빈 문단 (Word 요구사항)
    trailing_p = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr><w:spacing w:before="0" w:after="0" w:line="120" w:lineRule="auto"/></w:pPr>'
        f'</w:p>'
    )

    # 마지막 문단 뒤에 표 삽입 (순서 보장)
    last_p = tc.findall(qn('w:p'))[-1] if tc.findall(qn('w:p')) else None
    if last_p is not None:
        last_p.addnext(tbl)
    else:
        tc.append(tbl)
    # 표 뒤에 빈 문단 추가 (Word 요구: 셀 내 마지막 요소는 문단이어야 함)
    tbl.addnext(trailing_p)


def create_main_table(doc, sections_content):
    """본문 대형 표 생성"""
    # 행 수 계산: 섹션 1~5는 각 1행, 섹션 6은 1행 = 총 6행
    row_count = 6
    table = doc.add_table(rows=row_count, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, TABLE_WIDTH_DXA)
    set_table_borders(table, outer_sz=8, outer_color="2B5686", inner_sz=2, inner_color="999999")

    label_width = 1780   # 약 18.5%
    content_width = TABLE_WIDTH_DXA - label_width

    section_keys = ["1", "2", "3", "4", "5", "6"]

    for row_idx, sec_key in enumerate(section_keys):
        label_cell = table.cell(row_idx, 0)
        content_cell = table.cell(row_idx, 1)

        set_cell_width(label_cell, label_width)
        set_cell_width(content_cell, content_width)

        # 라벨 셀 스타일
        set_cell_shading(label_cell, COLOR_LABEL_BG)
        set_cell_vertical_align(label_cell, "center")
        set_cell_margins(label_cell, top=57, bottom=57, left=57, right=57)

        title, subtitle = SECTION_LABELS[sec_key]
        p = label_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, line=260)
        # 섹션 1은 "정의 및 중요성" 앞에 번호 없음 → 원본에서는 있음
        if sec_key == "1":
            display_title = "정의 및 중요성"
        else:
            display_title = title
        add_run(p, display_title, bold=True, size=10, spacing=-10)

        # 빈 줄
        p_space = label_cell.add_paragraph()
        set_paragraph_spacing(p_space, line=100)
        add_run(p_space, "", size=5)

        # 부제
        p_sub = label_cell.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p_sub, line=260)
        add_run(p_sub, subtitle, size=10, spacing=-10)

        # 내용 셀
        set_cell_margins(content_cell, top=57, bottom=57, left=113, right=113)
        if sec_key in sections_content:
            add_content_paragraphs(content_cell, sections_content[sec_key])
        else:
            p = content_cell.paragraphs[0]
            set_paragraph_spacing(p, line=260)

    return table


# ────────────────────────────────────────────
# Markdown 파서
# ────────────────────────────────────────────
def parse_md_table_rows(lines, start_idx):
    """Markdown 표 파싱 → 행 리스트"""
    rows = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith('|'):
            break
        if re.match(r'^\|[\s\-:|]+\|$', line):
            i += 1
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)
        i += 1
    return rows, i


def parse_markdown(md_text):
    """MD 텍스트를 파싱하여 구조화된 데이터 반환"""
    lines = md_text.split('\n')
    result = {
        'title_kr': '',
        'title_en': '',
        'mission': '',
        'affiliation': '',
        'name': '',
        'contact': '',
        'privacy': 'O',
        'sections': {}  # "1" ~ "6"
    }

    current_section = None
    section_lines = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 빈 줄
        if not stripped:
            if current_section and section_lines:
                # 본문에서 빈 줄은 무시 (문단 구분은 MD heading으로)
                pass
            i += 1
            continue

        # 수평선
        if stripped == '---':
            i += 1
            continue

        # H1: 문서 제목 (스킵)
        if stripped.startswith('# ') and not stripped.startswith('## '):
            i += 1
            continue

        # H2: 섹션 감지
        if stripped.startswith('## '):
            section_text = stripped[3:].strip()

            # 도전적 문제 표
            if '도전적 문제' in section_text and '정의서' not in section_text:
                # 다음에 표가 올 것
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('|'):
                    i += 1
                if i < len(lines):
                    rows, i = parse_md_table_rows(lines, i)
                    for row in rows:
                        if len(row) >= 2:
                            if '국문' in row[0]:
                                result['title_kr'] = row[1]
                            elif '영문' in row[0]:
                                result['title_en'] = row[1]
                continue

            # 미션
            if '미션' in section_text:
                i += 1
                while i < len(lines):
                    s = lines[i].strip()
                    if s and not s.startswith('#'):
                        result['mission'] = s
                        i += 1
                        break
                    elif s.startswith('#'):
                        break
                    i += 1
                continue

            # 제안자 인적사항
            if '제안자' in section_text or '인적사항' in section_text:
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('|'):
                    i += 1
                if i < len(lines):
                    rows, i = parse_md_table_rows(lines, i)
                    for row in rows:
                        if len(row) >= 2:
                            if '소속' in row[0]:
                                result['affiliation'] = row[1]
                            elif '성함' in row[0]:
                                result['name'] = row[1]
                            elif '연락처' in row[0]:
                                result['contact'] = row[1]
                            elif '개인정보' in row[0]:
                                result['privacy'] = row[1]
                continue

            # 본문 섹션 (1~6)
            sec_match = re.match(r'^(\d+)\.', section_text)
            if sec_match:
                # 이전 섹션 저장
                if current_section and section_lines:
                    result['sections'][current_section] = section_lines
                current_section = sec_match.group(1)
                section_lines = []
                i += 1
                continue

        # H3: 소제목
        if stripped.startswith('### '):
            if current_section:
                sub_text = stripped[4:].strip()
                section_lines.append({'type': 'heading3', 'text': sub_text})
            i += 1
            continue

        # H4: 소소제목
        if stripped.startswith('#### '):
            if current_section:
                sub_text = stripped[5:].strip()
                section_lines.append({'type': 'heading4', 'text': sub_text})
            i += 1
            continue

        # 표 감지
        if stripped.startswith('|') and i + 1 < len(lines) and lines[i + 1].strip().startswith('|'):
            if current_section:
                rows, i = parse_md_table_rows(lines, i)
                section_lines.append({'type': 'table', 'text': '', 'rows': rows})
            else:
                _, i = parse_md_table_rows(lines, i)
            continue

        # 불릿 리스트
        if stripped.startswith('- '):
            if current_section:
                section_lines.append({'type': 'bullet', 'text': stripped[2:].strip()})
            i += 1
            continue

        # 번호 리스트
        num_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if num_match and current_section:
            section_lines.append({'type': 'numbered', 'text': stripped})
            i += 1
            continue

        # 참고자료 섹션의 번호 항목 ([1], [2] ...)
        ref_match = re.match(r'^\[(\d+)\]\s+(.+)', stripped)
        if ref_match and current_section == "6":
            section_lines.append({'type': 'reference', 'text': stripped})
            i += 1
            continue

        # 일반 본문: 여러 줄 합치기
        if current_section:
            para_lines = [stripped]
            j = i + 1
            while j < len(lines):
                next_line = lines[j].strip()
                if (not next_line or
                    next_line.startswith('#') or
                    next_line.startswith('|') or
                    next_line.startswith('- ') or
                    re.match(r'^\d+\.\s+', next_line) or
                    re.match(r'^\[\d+\]', next_line) or
                    next_line == '---'):
                    break
                para_lines.append(next_line)
                j += 1
            full_text = ' '.join(para_lines)
            section_lines.append({'type': 'paragraph', 'text': full_text})
            i = j
            continue

        i += 1

    # 마지막 섹션 저장
    if current_section and section_lines:
        result['sections'][current_section] = section_lines

    return result


def add_inner_table(doc, cell, rows_data):
    """셀 안에 실제 Word 중첩 표를 삽입"""
    create_nested_table_in_cell(cell, rows_data)


def add_content_paragraphs_v2(cell, content_lines):
    """셀에 내용 추가 (표 포함)"""
    first = True
    for line_data in content_lines:
        line_type = line_data['type']
        text = line_data['text']

        if line_type == 'table':
            create_nested_table_in_cell(cell, line_data.get('rows', []))
            first = False  # 표 삽입 후 trailing_p가 추가되었으므로
            continue

        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()

        if line_type == 'heading3':
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(p, before=100, after=40, line=286)
            add_formatted_text(p, text, base_bold=True, base_size=10, spacing=-10)
        elif line_type == 'heading4':
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(p, before=60, after=30, line=286)
            add_formatted_text(p, text, base_bold=True, base_size=10, spacing=-10)
        elif line_type == 'bullet':
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(p, before=20, after=20, line=286)
            pPr = p._p.get_or_add_pPr()
            ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="284" w:hanging="142"/>')
            pPr.append(ind)
            add_run(p, "- ", size=10, spacing=-10)
            add_formatted_text(p, text, base_size=10, spacing=-10)
        elif line_type == 'numbered':
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(p, before=20, after=20, line=286)
            pPr = p._p.get_or_add_pPr()
            ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="284" w:hanging="284"/>')
            pPr.append(ind)
            add_formatted_text(p, text, base_size=10, spacing=-10)
        elif line_type == 'reference':
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(p, before=20, after=20, line=260)
            pPr = p._p.get_or_add_pPr()
            ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="284" w:hanging="284"/>')
            pPr.append(ind)
            add_formatted_text(p, text, base_size=10, spacing=-10)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(p, before=30, after=30, line=286)
            add_formatted_text(p, text, base_size=10, spacing=-10)

    if first:
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, line=260)


def convert_md_to_docx_form2(md_path, docx_path, header_text=None, footer_text=None):
    """서식2 도전적 문제 정의서: MD → DOCX 변환"""
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    data = parse_markdown(md_text)
    doc = Document()

    # ── 페이지 설정 ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)

    # ── 기본 스타일 ──
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # ── 머리글/바닥글 ──
    if header_text:
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = hp.add_run(header_text)
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    if footer_text:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run(footer_text)
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # ── 문서 제목 ──
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=0, line=120)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p_title, before=80, after=120)
    add_run(p_title, "도전적 문제 정의서", bold=True, size=16, font_name='맑은 고딕')

    # ── 표1: 도전적 문제 ──
    create_problem_table(doc, data['title_kr'], data['title_en'])

    # 표 간 간격
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=0, line=60)

    # ── 표2: 미션 + 제안자 ──
    create_info_table(doc, data['mission'], data['affiliation'], data['name'], data['contact'], data['privacy'])

    # 표 간 간격
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=0, line=60)

    # ── 표3: 본문 대형 표 ──
    row_count = 6
    table = doc.add_table(rows=row_count, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, TABLE_WIDTH_DXA)
    set_table_borders(table, outer_sz=8, outer_color="2B5686", inner_sz=2, inner_color="999999")

    label_width = 1780
    content_width = TABLE_WIDTH_DXA - label_width

    section_keys = ["1", "2", "3", "4", "5", "6"]
    display_labels = {
        "1": ("정의 및 중요성", "무엇이 문제인가?"),
        "2": ("2. 배경 및 목적", "왜 이 문제에 대한\n해결을 시도하려고\n하는가?"),
        "3": ("3. 목표 및 방법", "기존의 방법과 한계는\n무엇이며, 시도하고자\n하는 방법은?"),
        "4": ("4. 예상 결과", "문제를 해결할 수 있을\n것이라는 신호는\n무엇인가?"),
        "5": ("5. 파급효과", "만약 성공한다면,\n그 영향은 무엇인가?"),
        "6": ("6. 참고자료", "각 질문에 대해 입증할\n수 있는 학술 자료,\n통계, 보고서, 실험/\n임상 데이터 등의\n참고자료"),
    }

    for row_idx, sec_key in enumerate(section_keys):
        label_cell = table.cell(row_idx, 0)
        content_cell = table.cell(row_idx, 1)

        set_cell_width(label_cell, label_width)
        set_cell_width(content_cell, content_width)

        # 라벨 셀
        set_cell_shading(label_cell, COLOR_LABEL_BG)
        set_cell_vertical_align(label_cell, "center")
        set_cell_margins(label_cell, top=57, bottom=57, left=57, right=57)

        title, subtitle = display_labels[sec_key]
        p = label_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, line=260)
        add_run(p, title, bold=True, size=10, spacing=-10)

        p_space = label_cell.add_paragraph()
        set_paragraph_spacing(p_space, line=100)
        add_run(p_space, "", size=5)

        # 부제 (줄바꿈 처리)
        for sub_line in subtitle.split('\n'):
            p_sub = label_cell.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p_sub, line=260)
            add_run(p_sub, sub_line, size=10, spacing=-10)

        # 내용 셀
        set_cell_margins(content_cell, top=57, bottom=57, left=113, right=113)
        if sec_key in data['sections']:
            add_content_paragraphs_v2(content_cell, data['sections'][sec_key])
        else:
            p = content_cell.paragraphs[0]
            set_paragraph_spacing(p, line=260)

    doc.save(docx_path)
    print(f"✓ 생성 완료: {docx_path}")


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python3 md_to_docx_form2.py <input.md> <output.docx> [header_text] [footer_text]")
        sys.exit(1)

    md_path = sys.argv[1]
    docx_path = sys.argv[2]
    header = sys.argv[3] if len(sys.argv) > 3 else "한국형 ARPA-H 프로젝트 추진단 | 도전적 문제 정의서"
    footer = sys.argv[4] if len(sys.argv) > 4 else ""

    convert_md_to_docx_form2(md_path, docx_path, header, footer)
