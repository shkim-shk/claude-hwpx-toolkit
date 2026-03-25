#!/usr/bin/env python3
"""
MD → DOCX 변환 스크립트
ARPA-H Concept Paper 서식 재현:
- 맑은 고딕 10.5pt, 줄간격 110%
- 섹션 제목: 11.5pt bold #1F4E79
- 소제목: 10.5pt bold #2E75B6
- 표: 헤더 #1F4E79 배경 흰색 글자, 라벨 #E8F0FE 배경
- 여백: 상하 1.8cm, 좌 2.54cm, 우 2.0cm
- 머리글/바닥글 포함
"""
import re
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy


# ── 색상 상수 ──
COLOR_SECTION = RGBColor(0x1F, 0x4E, 0x79)   # 진한 남색
COLOR_SUBSECTION = RGBColor(0x2E, 0x75, 0xB6) # 파란색
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BG_HEADER = "1F4E79"
BG_LABEL = "E8F0FE"


def set_cell_shading(cell, hex_color):
    """셀 배경색 설정"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, sz=4, color="999999"):
    """셀 테두리 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def set_paragraph_spacing(para, before=None, after=None, line=None):
    """단락 간격 설정"""
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = parse_xml(f'<w:spacing {nsdecls("w")}/>')
        pPr.append(spacing)
    if before is not None:
        spacing.set(qn('w:before'), str(before))
    if after is not None:
        spacing.set(qn('w:after'), str(after))
    if line is not None:
        spacing.set(qn('w:line'), str(line))
        spacing.set(qn('w:lineRule'), 'auto')


def add_run(para, text, bold=False, italic=False, size=None, color=None, font_name=None):
    """서식이 적용된 run 추가"""
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return run


def add_formatted_text(para, text, base_bold=False, base_size=None, base_color=None):
    """Markdown 인라인 서식(**bold**, *italic*)을 파싱하여 run으로 추가"""
    # **bold** 와 *italic* 파싱
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))'
    for match in re.finditer(pattern, text):
        if match.group(2):  # ***bold italic***
            add_run(para, match.group(2), bold=True, italic=True, size=base_size, color=base_color)
        elif match.group(3):  # **bold**
            add_run(para, match.group(3), bold=True, size=base_size, color=base_color)
        elif match.group(4):  # *italic*
            add_run(para, match.group(4), bold=base_bold, italic=True, size=base_size, color=base_color)
        elif match.group(5):  # plain text
            add_run(para, match.group(5), bold=base_bold, size=base_size, color=base_color)


def calc_cell_spacing(text, is_header=False):
    """셀 내용 길이에 따라 줄간격/셀패딩을 동적으로 결정"""
    text_len = len(text.strip())
    if is_header:
        return {'before': 30, 'after': 30, 'line': 240}  # 헤더는 컴팩트
    elif text_len <= 15:
        # 짧은 내용 (단어 1~2개): 타이트하게
        return {'before': 20, 'after': 20, 'line': 240}
    elif text_len <= 50:
        # 중간 내용: 약간 여유
        return {'before': 30, 'after': 30, 'line': 260}
    elif text_len <= 120:
        # 긴 내용: 읽기 편하게
        return {'before': 40, 'after': 40, 'line': 280}
    else:
        # 매우 긴 내용 (문장 수준): 가독성 최대
        return {'before': 50, 'after': 50, 'line': 300}


def create_table(doc, rows_data, has_header=True):
    """표 생성 - ARPA-H 스타일, 내용 기반 동적 줄간격"""
    if not rows_data:
        return
    n_cols = len(rows_data[0])
    table = doc.add_table(rows=len(rows_data), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 표 너비 설정
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="9026" w:type="dxa"/>')
    tblPr.append(tblW)

    # 외곽 테두리
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # 열 너비 자동 계산: 각 열의 최대 텍스트 길이 기반
    col_max_lens = []
    for ci in range(n_cols):
        # 한글은 2, 영문/숫자는 1로 가중치 계산
        max_len = 0
        for row_data in rows_data:
            if ci < len(row_data):
                text = row_data[ci].strip()
                weighted = sum(2 if ord(c) > 0x7F else 1 for c in text)
                if weighted > max_len:
                    max_len = weighted
        # 최소 너비 보장 (헤더 텍스트 정도)
        col_max_lens.append(max(max_len, 4))

    total_weight = sum(col_max_lens)
    # 최소 비율 10% 보장, 나머지를 비례 배분
    min_ratio = 0.10
    col_widths = []
    remaining = 1.0 - min_ratio * n_cols
    if remaining < 0:
        remaining = 0
        min_ratio = 1.0 / n_cols
    for ci in range(n_cols):
        ratio = min_ratio + remaining * (col_max_lens[ci] / total_weight)
        col_widths.append(int(9026 * ratio))
    # 보정: 합계를 9026에 맞춤
    diff = 9026 - sum(col_widths)
    col_widths[-1] += diff

    for row_obj in table.rows:
        for ci, width in enumerate(col_widths):
            tc = row_obj.cells[ci]._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width}" w:type="dxa"/>')
            tcPr.append(tcW)

    for i, row_data in enumerate(rows_data):
        row = table.rows[i]
        # 행 내 가장 긴 셀 기준으로 줄간격 결정 (행 단위 통일)
        max_len_text = max(row_data, key=lambda t: len(t.strip()))
        is_hdr = (i == 0 and has_header)
        row_spacing = calc_cell_spacing(max_len_text, is_header=is_hdr)

        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(para,
                                  before=row_spacing['before'],
                                  after=row_spacing['after'],
                                  line=row_spacing['line'])

            clean_text = cell_text.strip()

            if i == 0 and has_header:
                set_cell_shading(cell, BG_HEADER)
                add_formatted_text(para, clean_text, base_bold=True, base_size=9.5, base_color=COLOR_WHITE)
            elif j == 0 and has_header:
                set_cell_shading(cell, BG_LABEL)
                add_formatted_text(para, clean_text, base_bold=True, base_size=9.5)
            else:
                add_formatted_text(para, clean_text, base_size=9.5)

    return table


def create_title_table(doc, rows_data):
    """서식1 제목 표 생성: 1열(제목 라벨)은 2행 병합, 2열(국문/영문), 3열(내용)"""
    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 표 너비
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="9026" w:type="dxa"/>')
    tblPr.append(tblW)

    # 외곽+내부 테두리
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # 열 너비 설정: 1열(라벨) ~2100, 2열(국/영) ~900, 3열(내용) ~6026
    for row_idx in range(2):
        row = table.rows[row_idx]
        for col_idx, width in enumerate([2100, 900, 6026]):
            tc = row.cells[col_idx]._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width}" w:type="dxa"/>')
            tcPr.append(tcW)

    # ── 1열: "제목: 해결하고 싶은 보건의료 분야 난제는?" — 2행 병합 ──
    label_text = rows_data[0][0].strip() if rows_data[0] else "제목"
    cell_a1 = table.cell(0, 0)
    cell_a2 = table.cell(1, 0)
    cell_a1.merge(cell_a2)  # rowSpan=2

    set_cell_shading(cell_a1, "E5E5E5")
    para = cell_a1.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(para, before=40, after=40, line=280)
    add_run(para, label_text, bold=True, size=10, color=COLOR_SECTION)

    # 수직 가운데 정렬
    tc_pr = cell_a1._tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tc_pr.append(vAlign)

    # ── 2열, 3열: 국문/영문 ──
    for row_idx in range(2):
        row = table.rows[row_idx]
        # 2열: 국문/영문 라벨
        label_cell = row.cells[1]
        set_cell_shading(label_cell, "E5E5E5")
        lp = label_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(lp, before=40, after=40, line=280)
        label = rows_data[row_idx][1].strip() if len(rows_data[row_idx]) > 1 else ""
        add_run(lp, label, bold=True, size=10)
        # 수직 가운데
        tc_pr2 = label_cell._tc.get_or_add_tcPr()
        vAlign2 = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
        tc_pr2.append(vAlign2)

        # 3열: 내용
        content_cell = row.cells[2]
        cp = content_cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(cp, before=40, after=40, line=280)
        content = rows_data[row_idx][2].strip() if len(rows_data[row_idx]) > 2 else ""
        add_formatted_text(cp, content, base_size=10)
        # 수직 가운데
        tc_pr3 = content_cell._tc.get_or_add_tcPr()
        vAlign3 = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
        tc_pr3.append(vAlign3)

    return table


def parse_md_table(lines, start_idx):
    """Markdown 표 파싱. (헤더행, 구분행, 데이터행들) 반환"""
    rows = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith('|'):
            break
        # 구분행(---|---) 건너뛰기
        if re.match(r'^\|[\s\-:|]+\|$', line):
            i += 1
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)
        i += 1
    return rows, i


def add_header_footer(doc, header_text, footer_text):
    """머리글/바닥글 설정"""
    section = doc.sections[0]
    # 머리글
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run(header_text)
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # 바닥글
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(footer_text)
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')


def convert_md_to_docx(md_path, docx_path, header_text=None, footer_text=None):
    """MD → DOCX 변환 (ARPA-H 서식)"""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    doc = Document()

    # ── 페이지 설정 ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    # ── 기본 스타일 설정 ──
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(2.5)
    pf.space_after = Pt(2.5)
    # 줄간격 110%
    pf.line_spacing_rule = None  # auto

    # ── 머리글/바닥글 ──
    if header_text:
        add_header_footer(doc, header_text, footer_text or "")

    # ── 본문 파싱 및 생성 ──
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 빈 줄
        if not stripped:
            i += 1
            continue

        # 수평선 (---)
        if stripped == '---':
            i += 1
            continue

        # Markdown 표 감지
        if stripped.startswith('|') and i + 1 < len(lines) and lines[i+1].strip().startswith('|'):
            rows, i = parse_md_table(lines, i)
            if rows:
                # 제목 표 감지: 첫 셀에 "제목" 포함
                is_title_table = any('제목' in rows[0][0] for _ in [1]) if rows else False
                if is_title_table and len(rows) >= 2:
                    create_title_table(doc, rows)
                else:
                    create_table(doc, rows, has_header=True)
                # 표 뒤 빈 단락
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=0, after=40)
            continue

        # H1: 문서 제목 — 13pt (본문+3)
        if stripped.startswith('# ') and not stripped.startswith('## '):
            title_text = stripped[2:].strip()
            # 제목 전 빈 줄
            spacer = doc.add_paragraph()
            set_paragraph_spacing(spacer, before=0, after=0, line=120)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, before=80, after=40)
            add_run(p, title_text, bold=True, size=13, color=COLOR_SECTION)
            i += 1
            continue

        # H2: 섹션 제목 (□) — 12pt (본문+2)
        if stripped.startswith('## '):
            section_text = stripped[3:].strip()
            # 섹션 전 빈 줄
            spacer = doc.add_paragraph()
            set_paragraph_spacing(spacer, before=0, after=0, line=160)
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=100, after=80)
            add_formatted_text(p, section_text, base_bold=True, base_size=12, base_color=COLOR_SECTION)
            i += 1
            continue

        # H3: 소제목 (❍) — 11pt (본문+1)
        if stripped.startswith('### '):
            sub_text = stripped[4:].strip()
            # 소제목 전 빈 줄
            spacer = doc.add_paragraph()
            set_paragraph_spacing(spacer, before=0, after=0, line=120)
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=60, after=50)
            add_formatted_text(p, sub_text, base_bold=True, base_size=11, base_color=COLOR_SUBSECTION)
            i += 1
            continue

        # 불릿 리스트
        if stripped.startswith('- '):
            list_text = stripped[2:].strip()
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=30, after=30, line=264)
            pPr = p._p.get_or_add_pPr()
            ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="360" w:hanging="180"/>')
            pPr.append(ind)
            add_run(p, "• ", bold=False, size=10)
            add_formatted_text(p, list_text, base_size=10)
            i += 1
            continue

        # 번호 리스트 (1. 2. 3. ...) — 각각 별도 단락
        num_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if num_match:
            num = num_match.group(1)
            list_text = num_match.group(2).strip()
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=30, after=30, line=264)
            pPr = p._p.get_or_add_pPr()
            ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="360" w:hanging="360"/>')
            pPr.append(ind)
            add_run(p, f"{num}. ", bold=False, size=10)
            add_formatted_text(p, list_text, base_size=10)
            i += 1
            continue

        # 일반 본문 단락
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=50, after=50, line=264)

        # 여러 줄을 하나의 단락으로 합치기 (빈 줄이나 특수 요소 전까지)
        para_lines = [stripped]
        j = i + 1
        while j < len(lines):
            next_line = lines[j].strip()
            # 단락 종료 조건
            if (not next_line or
                next_line == '---' or
                next_line.startswith('#') or
                next_line.startswith('|') or
                next_line.startswith('- ') or
                re.match(r'^\d+\.\s+', next_line) or
                next_line.startswith('*Go/No-Go*') or
                next_line.startswith('*리스크')):
                break
            para_lines.append(next_line)
            j += 1

        full_text = ' '.join(para_lines)
        add_formatted_text(p, full_text, base_size=10)
        i = j
        continue

    doc.save(docx_path)
    print(f"생성 완료: {docx_path}")


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python3 md_to_docx.py <input.md> <output.docx> [header_text] [footer_text]")
        sys.exit(1)

    md_path = sys.argv[1]
    docx_path = sys.argv[2]
    header = sys.argv[3] if len(sys.argv) > 3 else "한국형 ARPA-H 프로젝트 추진단 | Concept Paper (예비제안서)"
    footer = sys.argv[4] if len(sys.argv) > 4 else "미정복질환(희귀) 극복분야"

    convert_md_to_docx(md_path, docx_path, header, footer)
